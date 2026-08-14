from __future__ import annotations

import io

import docx
import pytest
from pypdf import PdfWriter

from django_app.rag.extract import UnsupportedFileTypeError, extract_text


def test_extract_text_from_txt():
    assert extract_text("notes.txt", b"hello world") == "hello world"


def test_extract_text_from_md():
    assert extract_text("notes.md", b"# Title\n\nBody") == "# Title\n\nBody"


def test_extract_text_unsupported_extension_raises():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text("archive.zip", b"whatever")


def test_extract_text_no_extension_raises():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text("noextension", b"whatever")


def test_extract_text_from_docx_round_trip():
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("doc.docx", buffer.getvalue())

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_text_from_pdf_does_not_crash():
    # A real text-bearing PDF fixture would need a rendering library this
    # project doesn't otherwise depend on — a blank page is enough to prove
    # PdfReader wiring (BytesIO in, str out) without one.
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    text = extract_text("doc.pdf", buffer.getvalue())

    assert isinstance(text, str)
